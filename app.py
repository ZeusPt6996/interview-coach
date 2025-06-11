import streamlit as st
st.set_page_config(page_title="InterviewCoach Pro", layout="wide")
import openai
import fitz  # PyMuPDF for PDF text extraction
import re
from docx import Document
from io import BytesIO

st.title("🎯 InterviewCoach Pro")
st.markdown("Your AI Mock Interviewer for High-Performance Roles")

# --- API KEY SETUP ---
import os
openai.api_key = st.secrets["openai_api_key"]

if "step" not in st.session_state:
    st.session_state.step = 1

# --- Step 1: Upload & Setup ---
if st.session_state.step == 1:
    st.header("🗕️ Step 1: Upload Job Description & Resume")
    jd_input = st.text_area("📌 Paste the Job Description:")
    cv_file = st.file_uploader("📄 Upload your resume (PDF only):", type=["pdf"])
    jd_q_count = st.number_input("🎯 How many questions from the Job Description?", min_value=0, max_value=10, value=2)
    cv_q_count = st.number_input("📚 How many questions from your Resume?", min_value=0, max_value=10, value=2)

    if jd_input and cv_file:
        cv_text = ""
        with fitz.open(stream=cv_file.read(), filetype="pdf") as doc:
            for page in doc:
                cv_text += page.get_text()
        st.session_state['jd'] = jd_input
        st.session_state['cv'] = cv_text
        st.session_state['jd_q_count'] = jd_q_count
        st.session_state['cv_q_count'] = cv_q_count
        with st.expander("🔍 Preview Extracted Resume Text"):
            st.text_area("Extracted CV Text:", value=cv_text, height=300)
        if st.button("➡️ Proceed to Step 2"):
            st.session_state.step = 2

# --- Step 2: Generate Questions ---
elif st.session_state.step == 2:
    st.header("❓ Step 2: Answer Tailored Mock Interview Questions")
    with st.expander("💡 Need help with STAR format?"):
        st.markdown("""
        **S – Situation:** Set the scene  
        **T – Task:** What you had to achieve  
        **A – Action:** Steps you took  
        **R – Result:** Outcome & impact (preferably with numbers)
        """)

    if 'questions' not in st.session_state:
        prompt = f"""
        You are a mock interview coach. Generate a total of {st.session_state['jd_q_count'] + st.session_state['cv_q_count']} behavioral interview questions.
        - First, generate {st.session_state['jd_q_count']} questions based strictly on the Job Description below.
        - Then, generate {st.session_state['cv_q_count']} questions based on the candidate's resume.
        Job Description:
        {st.session_state['jd']}
        Resume:
        {st.session_state['cv']}
        Format: Numbered list.
        """
        try:
            with st.spinner("Generating questions..."):
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}]
                )
            raw = response.choices[0].message.content.strip().split('\n')
            filtered = []
            for line in raw:
                if re.match(r"^\d+\.\s.+", line):
                    parts = line.split(". ", 1)
                    if len(parts) == 2:
                        q = parts[1].strip()
                        if q:
                            filtered.append(q)
            total_qs = st.session_state['jd_q_count'] + st.session_state['cv_q_count']
            st.session_state['questions'] = filtered[:total_qs]
        except Exception as e:
            st.error(f"Question generation failed: {e}")
            st.stop()

    all_answered = True
    for i, q in enumerate(st.session_state['questions']):
        st.markdown(f"**Q{i+1}:** {q}")
        answer = st.text_area(f"Your Answer to Q{i+1}", key=f"answer_{i}")
        if not answer.strip():
            all_answered = False

    if all_answered and st.button("➡️ Proceed to Step 3"):
        st.session_state.step = 3

# --- Step 3: Get Feedback ---
elif st.session_state.step == 3:
    st.header("🧠 Step 3: Receive Feedback & Suggested Improvements")
    final_score_total = 0
    valid_scores = 0
    feedback_export = []

    for i, q in enumerate(st.session_state['questions']):
        answer = st.session_state.get(f"answer_{i}", "")

        feedback_prompt = f'''
You are a high-stakes mock interview coach for elite consulting, marketing, sales, and product roles. Evaluate critically using the STAR method with McKinsey-level scrutiny. Be rigorous.

## Step 1: STAR Breakdown
Break the answer into Situation, Task, Action, Result. 
- Identify weak verbs, filler content, and gaps in logic.
- Penalize generic or unverifiable claims.
- Praise tangible impact and confident storytelling.

## Step 2: Final Score (out of 10)
- 9–10: Elite answer; confident, structured, quantifiable.
- 7–8: Solid but could use clearer impact or tighter phrasing.
- 5–6: Adequate structure, but weak delivery or result.
- <5: Confusing or structurally broken.

End like: Score: 7.5/10 – Strong action but impact unclear.

## Step 3: Resume-Based Enhancement
Cross-reference with CV. Suggest better framing or context from resume items.

Question: {q}  
Answer: {answer}  
Resume: {st.session_state['cv']}
'''

        rewrite_prompt = f"""
Rewrite this answer in better STAR format:
- Sharpen the task
- Add results
- Make it business-relevant and crisp

Question: {q}
Original Answer: {answer}
"""

        try:
            with st.spinner(f"Generating feedback for Q{i+1}..."):
                feedback_response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": feedback_prompt}]
                )
                feedback = feedback_response.choices[0].message.content.strip()

                rewrite_response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": rewrite_prompt}]
                )
                rewrite = rewrite_response.choices[0].message.content.strip()

            with st.expander(f"📝 Feedback for Q{i+1}"):
                st.markdown(f"**🔹 Question:** {q}")

                with st.container():
    st.markdown("### 💬 Original Answer")
    import textwrap
    wrapped_answer = textwrap.fill(answer, width=100)
    st.markdown(wrapped_answer)

                with st.container():
                    st.markdown("### 📌 Feedback")
                    import textwrap
wrapped_feedback = textwrap.fill(feedback, width=100)
st.markdown(wrapped_feedback, unsafe_allow_html=True)

                with st.container():
                    st.markdown("### ✍️ Suggested Rewrite")
                    st.text_area("Enhanced Answer", value=rewrite, height=200)

            feedback_export.append(f"Q{i+1}: {q}\n\nFEEDBACK:\n{feedback}\n\nSUGGESTED REWRITE:\n{rewrite}\n")

            score_match = re.search(r"Score:\s*(\d+(\.\d+)?)\s*/\s*10", feedback, re.IGNORECASE)
            if score_match:
                final_score_total += float(score_match.group(1))
                valid_scores += 1
        except Exception as e:
            st.error(f"Error in generating feedback or rewrite: {e}")

    if feedback_export:
        st.subheader("📊 Step 4: Summary Report & Final Recommendation")
        if valid_scores:
            avg_score = final_score_total / valid_scores
            st.success(f"✅ Average STAR Answer Score: {avg_score:.1f}/10")
        else:
            avg_score = 0
            st.warning("⚠️ Average Score: Not available")

        fit_prompt = f"""
You are a senior hiring manager trained in candidate evaluation. Using the Job Description, Resume, and the candidate's STAR-format interview responses, assess how well this candidate fits the role.

Rate on these 5 dimensions (1–10 each):
- Technical Tools: List specific tools used (SQL, Tableau, etc.), assess alignment with JD.
- Domain Experience: Match real work done with what the role demands.
- Problem Solving: Reference concrete examples and business impact from answers.
- Communication: Judge structure, clarity, and precision in written answers.
- Leadership/Initiative: Use resume and answers to assess proactive behavior.

For each dimension, first give the score (e.g., "Technical Tools: 8") on a line, then a line break, then explain why.

Conclude with a total out of 50, percentage out of 100, and a final one-line hiring recommendation.

Job Description:
{st.session_state['jd']}

Resume:
{st.session_state['cv']}

Interview Answers:
{[st.session_state.get(f"answer_{i}", '') for i in range(len(st.session_state['questions']))]}
"""

        try:
            with st.spinner("Evaluating candidate fit score..."):
                fit_response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": fit_prompt}]
                )
                fit_result = fit_response.choices[0].message.content.strip()
                st.info(f"📌 Fit Score Summary (fit between JD and CV):\n\n{fit_result}")

                recommendation = "✅ Recommendation: Proceed to interview" if avg_score >= 8 else ("🚫 Recommendation: Not ready for interview" if avg_score < 7 else "🚫 Recommendation: Practice more before the interview but you are close")
                st.markdown(recommendation)
        except Exception as e:
            st.error(f"Error generating fit score: {e}")

        report_doc = Document()
        report_doc.add_heading("InterviewCoach Pro - Feedback Report", level=1)
        for i, item in enumerate(feedback_export):
            report_doc.add_page_break()
            report_doc.add_heading(f"Question {i+1}", level=2)
            sections = item.split("\n\n")
            for section in sections:
                report_doc.add_paragraph(section.strip())

        word_file = BytesIO()
        report_doc.save(word_file)
        word_file.seek(0)

        st.download_button(
            label="🗕️ Download Full Feedback Report (.docx)",
            data=word_file,
            file_name="InterviewCoach_Feedback.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
